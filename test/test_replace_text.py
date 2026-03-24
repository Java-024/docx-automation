from docx import Document

def replace_text(file_path, output_path, search_text, replace_text):
    doc = Document(file_path)

    # Замена в абзацах
    for paragraph in doc.paragraphs:
        if search_text in paragraph.text:
            paragraph.text = paragraph.text.replace(search_text, replace_text)

    # Замена в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if search_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(search_text, replace_text)

    doc.save(output_path)
    print(f"Файл сохранен: {output_path}")

# Использование
replace_text('document.docx', 'result.docx', 'юного', 'проверка сохранения форматирования текста')
