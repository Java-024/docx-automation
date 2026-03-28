from docx import Document


def find_structures(file_path):
    """
    Поиск структур в документе Word, начинающихся с # и одного из символов: ", *, =, \
    и заканчивающихся тем же символом. & используется для объединения структур.
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"Ошибка при открытии файла: {e}")
        return []

    # Собираем весь текст из документа
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text.append(paragraph.text)

    text = '\n'.join(full_text)

    # Специальные символы для начала и конца структуры
    special_chars = {'"', '*', '=', '\\'}

    # Функция для обработки & как объединителя структур
    def process_ampersand(text):
        """Обработка & как объединителя структур"""
        results = []
        i = 0

        while i < len(text):
            if text[i] == '#':
                if i + 1 < len(text) and text[i + 1] in special_chars:
                    start_char = text[i + 1]
                    structure = ['#', start_char]
                    j = i + 2

                    while j < len(text):
                        if text[j] == '&':
                            # & встречается - продолжаем, не закрывая структуру
                            structure.append(text[j])
                            j += 1
                        elif text[j] == start_char:
                            structure.append(start_char)
                            results.append(''.join(structure))
                            i = j
                            break
                        else:
                            structure.append(text[j])
                            j += 1
                    else:
                        i += 1
                else:
                    i += 1
            else:
                i += 1

        return results

    # Получаем структуры
    structures = process_ampersand(text)

    # Выводим результаты
    if structures:
        print("Найденные структуры:")
        for struct in structures:
            print(struct)
    else:
        print("Структуры не найдены.")

    return structures


if __name__ == "__main__":
    # Здесь укажи путь к твоему файлу
    file_path = "test_search.docx"  # <-- ЗДЕСЬ УКАЖИ ПУТЬ

    find_structures(file_path)